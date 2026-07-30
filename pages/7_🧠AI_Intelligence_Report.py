import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from docx import Document
from utils.data_loader import load_data

st.set_page_config(page_title="AI Intelligence Report", page_icon="🧠", layout="wide")
st.title("🧠 AI Intelligence Report")
st.write("Auto-generated summary analysis of global terrorism trends from the dataset.")

df = load_data()

total_attacks = len(df)
total_deaths = int(df["nkill"].sum())
total_injured = int(df["nwound"].sum())
top_country = df["country_txt"].value_counts().idxmax()
top_group = df["gname"].value_counts().idxmax()
top_attack_type = df["attacktype1_txt"].value_counts().idxmax()
year_range = f"{int(df['iyear'].min())} - {int(df['iyear'].max())}"

c1, c2, c3, c4 = st.columns(4)
c1.metric("Total Incidents", f"{total_attacks:,}")
c2.metric("Total Fatalities", f"{total_deaths:,}")
c3.metric("Total Injured", f"{total_injured:,}")
c4.metric("Years Covered", year_range)

st.markdown("---")

st.subheader("Key Insights")
st.markdown(f"""
- **Most affected country:** {top_country}
- **Most active group on record:** {top_group}
- **Most common attack type:** {top_attack_type}
- Dataset spans **{year_range}**, covering **{total_attacks:,} recorded incidents**.
""")

col1, col2 = st.columns(2)

with col1:
    top10_countries = df["country_txt"].value_counts().head(10).reset_index()
    top10_countries.columns = ["Country", "Attacks"]
    fig1 = px.bar(top10_countries, x="Attacks", y="Country", orientation="h",
                  title="Top 10 Most Affected Countries", color="Attacks",
                  color_continuous_scale="Reds")
    fig1.update_layout(yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig1, use_container_width=True)

with col2:
    attack_types = df["attacktype1_txt"].value_counts().reset_index()
    attack_types.columns = ["Attack Type", "Count"]
    fig2 = px.pie(attack_types, names="Attack Type", values="Count",
                  title="Attack Type Distribution", hole=0.4)
    st.plotly_chart(fig2, use_container_width=True)

trend = df.groupby("iyear").size().reset_index(name="Attacks")
fig3 = px.line(trend, x="iyear", y="Attacks", title="Attacks Over Time", markers=True)
st.plotly_chart(fig3, use_container_width=True)

st.markdown("---")


def generate_pdf():
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 16)
    c.drawString(2 * cm, height - 2 * cm, "AI Military Intelligence Report")

    c.setFont("Helvetica", 11)
    lines = [
        f"Total Incidents: {total_attacks:,}",
        f"Total Fatalities: {total_deaths:,}",
        f"Total Injured: {total_injured:,}",
        f"Years Covered: {year_range}",
        "",
        f"Most Affected Country: {top_country}",
        f"Most Active Group: {top_group}",
        f"Most Common Attack Type: {top_attack_type}",
    ]

    y = height - 3 * cm
    for line in lines:
        c.drawString(2 * cm, y, line)
        y -= 0.8 * cm

    c.save()
    buffer.seek(0)
    return buffer


def generate_docx():
    doc = Document()
    doc.add_heading("AI Military Intelligence Report", level=1)

    doc.add_paragraph(f"Total Incidents: {total_attacks:,}")
    doc.add_paragraph(f"Total Fatalities: {total_deaths:,}")
    doc.add_paragraph(f"Total Injured: {total_injured:,}")
    doc.add_paragraph(f"Years Covered: {year_range}")

    doc.add_heading("Key Insights", level=2)
    doc.add_paragraph(f"Most Affected Country: {top_country}")
    doc.add_paragraph(f"Most Active Group: {top_group}")
    doc.add_paragraph(f"Most Common Attack Type: {top_attack_type}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


col1, col2 = st.columns(2)
with col1:
    st.download_button(
        "📄 Download PDF Report",
        data=generate_pdf(),
        file_name="AI_Intelligence_Report.pdf",
        mime="application/pdf"
    )
with col2:
    st.download_button(
        "📝 Download Word Report",
        data=generate_docx(),
        file_name="AI_Intelligence_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )